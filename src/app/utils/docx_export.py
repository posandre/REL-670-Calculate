from __future__ import annotations

from base64 import b64decode
from io import BytesIO
from pathlib import Path
from xml.etree import ElementTree

from docx import Document
from docx.shared import Inches, RGBColor


def export_html_to_docx(html: str, path: Path) -> None:
    document = Document()
    root = ElementTree.fromstring(f"<root>{html}</root>")
    for child in root:
        _append_block(document, child)
    document.save(path)


def _append_block(document: Document, element: ElementTree.Element) -> None:
    tag = _tag(element)
    if tag == "h2":
        _append_heading(document, element, level=1)
    elif tag == "h3":
        _append_heading(document, element, level=2)
    elif tag == "p":
        _append_paragraph(document, element)
    elif tag in {"ol", "ul"}:
        _append_list(document, element, numbered=tag == "ol")
    elif tag == "table":
        _append_table(document, element)


def _append_heading(document: Document, element: ElementTree.Element, *, level: int) -> None:
    paragraph = document.add_heading(level=level)
    _append_inline(paragraph, element)


def _append_paragraph(
    document: Document,
    element: ElementTree.Element,
    *,
    style: str | None = None,
) -> None:
    paragraph = document.add_paragraph(style=style)
    _append_inline(paragraph, element)


def _append_list(document: Document, element: ElementTree.Element, *, numbered: bool) -> None:
    style = "List Number" if numbered else "List Bullet"
    for item in element:
        if _tag(item) != "li":
            continue
        first_text = _direct_text(item)
        wrote_leading_item = False
        if first_text:
            paragraph = document.add_paragraph(style=style)
            run = paragraph.add_run(first_text)
            _force_black(run)
            wrote_leading_item = True
        first_paragraph = not wrote_leading_item
        for child in item:
            child_tag = _tag(child)
            if child_tag == "p":
                _append_paragraph(document, child, style=style if first_paragraph else None)
                first_paragraph = False
            elif child_tag in {"ol", "ul"}:
                _append_list(document, child, numbered=child_tag == "ol")
            elif child_tag == "table":
                _append_table(document, child)


def _append_table(document: Document, element: ElementTree.Element) -> None:
    row_elements = [row for row in element if _tag(row) == "tr"]
    if not row_elements:
        return
    column_count = max(
        len([cell for cell in row if _tag(cell) in {"th", "td"}])
        for row in row_elements
    )
    table = document.add_table(rows=len(row_elements), cols=column_count)
    table.style = "Table Grid"
    for row_index, row_element in enumerate(row_elements):
        cells = [cell for cell in row_element if _tag(cell) in {"th", "td"}]
        for column_index, cell_element in enumerate(cells):
            cell = table.cell(row_index, column_index)
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(_text(cell_element))
            run.bold = _tag(cell_element) == "th"
            _force_black(run)


def _append_inline(
    paragraph,
    element: ElementTree.Element,
    *,
    bold: bool = False,
    subscript: bool = False,
) -> None:
    if element.text:
        run = paragraph.add_run(element.text)
        run.bold = bold
        run.font.subscript = subscript
        _force_black(run)
    for child in element:
        child_tag = _tag(child)
        if child_tag == "img":
            _append_image(paragraph, child)
        else:
            child_bold = bold or child_tag in {"b", "strong", "th"}
            _append_inline(
                paragraph,
                child,
                bold=child_bold,
                subscript=subscript or child_tag == "sub",
            )
        if child.tail:
            run = paragraph.add_run(child.tail)
            run.bold = bold
            run.font.subscript = subscript
            _force_black(run)


def _append_image(paragraph, element: ElementTree.Element) -> None:
    source = element.attrib.get("src", "")
    if source.startswith("data:image/png;base64,"):
        image_data = BytesIO(b64decode(source.split(",", 1)[1]))
        run = paragraph.add_run()
        run.add_picture(image_data, width=Inches(6.5))


def _force_black(run) -> None:
    run.font.color.rgb = RGBColor(0, 0, 0)


def _tag(element: ElementTree.Element) -> str:
    return element.tag.lower()


def _text(element: ElementTree.Element) -> str:
    return "".join(element.itertext()).strip()


def _direct_text(element: ElementTree.Element) -> str:
    return (element.text or "").strip()
